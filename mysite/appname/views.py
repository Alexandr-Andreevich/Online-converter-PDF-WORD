from pyexpat import model
from django.forms import ModelChoiceField
from django.shortcuts import render, redirect
from django.http import HttpResponse

from .forms import *
from django.contrib.auth.models import User
from django.contrib import auth
import docx2pdf
import pdfplumber
from docx import Document


def page_for_ghost(request):  # Страница гостя (главаня)
    return render(request, "page_for_ghost.html")


def register(request):  # Регистрация
    if request.method == "POST":
        username = request.POST["username"]
        password1 = request.POST["password1"]
        password2 = request.POST["password2"]
        if password1 == password2:
            if User.objects.filter(username=username).exists():
                return render(
                    request,
                    "register.html",
                    {"error": "Такой пользователь уже существует!"},
                )
            else:
                user = User.objects.create_user(username=username, password=password1)
                user.save()
                return redirect("login")
        else:
            return render(request, "register.html", {"error": "Пароли не совпадают!"})
    else:
        return render(request, "register.html")


def login(request):  # Вход
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = auth.authenticate(username=username, password=password)
        if user is not None:
            auth.login(request, user)
            return redirect("choise_format")
        else:
            return render(request, "login.html", {"error": "Введены неверные данные!"})
    else:
        return render(request, "login.html")


def choise_format(request):  # Страница пользователя (выбор формата конвертации)
    return render(request, "choise_format.html")


def change_to_word(request):  # Страница конвертации в WORD формат
    if request.method == "POST":
        pdf_to_word(request.FILES["load_file"])
    return render(request, "change_to_word.html")


def change_to_pdf(request):  # Страница конвертации в PDF формат
    if request.method == "POST":
        word_to_pdf(request.FILES["load_file"])
        return redirect("view_pdf_file")
    return render(request, "change_to_pdf.html")


def view_pdf_file(request):  # Страница просмотра PDF файла
    return render(request, "view_pdf_file.html")


def word_to_pdf(f):  # Функция конвертации WORD в PDF
    docx2pdf.convert(f"{f.name}", f"appname/static/uploads_file/Esenin.pdf")


def pdf_to_word(f):  # Функция конвертации PDF в WORD
    pdf_file = pdfplumber.open(f"appname/static/uploads_file/Esenin.pdf")
    doc = Document()
    for page in pdf_file.pages:
        text = page.extract_text()
        doc.add_paragraph(text)
    doc.save(f"appname/static/uploads_file/Esenin.docx")

        # или можно так
        # Перевод файла в WORD

# from docx import Document 
# import PyPDF2


# def pdf_to_docx(input_pdf_path, output_docx_path):
#     """
#     Функция преобразования из PDF в DOCX
#     Принимает путь PDF файла и путь для сохранения DOCX файла
#     """
#     pdf = PyPDF2.PdfReader(input_pdf_path)
#     document = Document()

#     for page_num in range(len(pdf.pages)):
#         page = pdf.pages[page_num]
#         text = page.extract_text()
#         document.add_paragraph(text)

#     document.save(output_docx_path)
#     print(f'Конвертация файла из PDF в DOCX завершена, файл сохранен по пути: {output_docx_path}')


# if __name__ == '__main__':
#     input_pdf_path = 'path/to/file.pdf'
#     output_docx_path = 'path/to/file.docx'
#     pdf_to_docx(input_pdf_path, output_docx_path)
    
# 1) from docx import Document: Эта строка импортирует класс Document из библиотеки docx, который используется для создания нового документа Word.

# 2) import PyPDF2: Эта строка импортирует библиотеку PyPDF2, которая используется для чтения PDF-файлов.

# 3) def pdf_to_docx(input_pdf_path, output_docx_path): Это определение функции pdf_to_docx, которая принимает два аргумента: путь к входному PDF-файлу и путь к выходному DOCX-файлу.

# 4) pdf = PyPDF2.PdfReader(input_pdf_path): Эта строка создает объект PdfReader, который считывает содержимое входного PDF-файла.

# 5) document = Document(): Эта строка создает новый документ Word.

# 6) for page_num in range(len(pdf.pages)): Этот цикл проходит по каждой странице в PDF-файле.

# 7) page = pdf.pages[page_num]: Эта строка получает текущую страницу из PDF-файла.

# 8) text = page.extract_text(): Эта строка извлекает текст из текущей страницы.

# 9) document.add_paragraph(text): Эта строка добавляет извлеченный текст в документ Word в виде нового абзаца.

# 10) document.save(output_docx_path): Эта строка сохраняет документ Word в указанном выходном файле.

# 11) print(f'Конвертация файла из PDF в DOCX завершена, файл сохранен по пути: {output_docx_path}'): Эта строка выводит сообщение, подтверждающее успешное завершение конвертации.

# 12) if __name__ == '__main__': Это условие гарантирует, что код будет выполнен только при прямом запуске этого скрипта, а не при его импорте как модуля.

# 13) input_pdf_path = 'path/to/file.pdf', output_docx_path = 'path/to/file.docx', pdf_to_docx(input_pdf_path, output_docx_path): Эти строки задают пути к входному и выходному файлам и вызывают функцию pdf_to_docx для выполнения конвертации.
# --------------------------------------------------------


# --------------------------------------------------------
        # Функция загрузки файла
    
# def upload_file(f): 
#     with open(f"uploads/{f.name}", "wb+") as destination:
#         for chunk in f.chunks():
#             destination.write(chunk)
#     print(f)
# --------------------------------------------------------


# --------------------------------------------------------
        # Перевод файла в PDF
    
# from fpdf import FPDF 

# pdf = FPDF()

# pdf.add_page()
# pdf.set_font('Arial', size=16)

# with open('Plain.txt', 'r') as f:
#     for i in f:
#         pdf.cell(200, 10, txt=i, ln=1, align='C')

# pdf.output('Plainpdf.pdf')

        # или

# doc = Document('Esenin.docx')
# full_text = []
# for p in doc.paragraphs:
#     full_text.append(p.text)
# full_text = '\n\n'.join(full_text)
# doc.save('Esenin.pdf')
# print(full_text)

        # или

# pdf = FPDF()

# pdf.add_page()
# pdf.add_font('DejaVu', '', 'font\DejaVuSansCondensed.ttf', uni=True)
# pdf.set_font("DejaVu", size=14)

# with open(f"{f.name}", "r") as f:
#     for i in f:
#         pdf.cell(200, 10, txt=i, ln=1, align="C")

# pdf.output(f"uploads/{f.name}.pdf")

        # или

# doc = aw.Document(f"{f.name}")
# doc.save(f"uploads/Esenin.pdf")
# --------------------------------------------------------


# --------------------------------------------------------
# Перевод файла в WORD
    
# def change_to_word():
#     pdf_file = pdfplumber.open('GK_JavaScrip.pdf')
#     doc = Document()

#     for page in pdf_file.pages:
#         text = page.extract_text()
#         doc.add_paragraph(text)

#     doc.save('Gk-JavaScript.doxc')
# --------------------------------------------------------


# doc = docx.Document(f)
# for paragraphs in doc.paragraphs:
#     return HttpResponse(paragraphs.text)